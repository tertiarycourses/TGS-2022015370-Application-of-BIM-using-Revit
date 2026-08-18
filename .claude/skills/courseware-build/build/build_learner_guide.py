#!/usr/bin/env python3
"""Generate the WSQ Application of BIM using Revit (TGS-2022015370) Learner Guide
as BOTH a Markdown mirror (LG-*.md at repo root) and a DOCX (courseware/LG-*.docx)
from one source, so they never diverge.

House format: cover page, Document Version Control Record, auto TOC, Arial 11pt
body, one section per lab (Objective · Goal · What you'll build · Step-by-step
with commands · Test it), plus setup, exam-prep and glossary. All content is
driven by course_data + the domain data files, keeping the LG 100% aligned with
the slide deck, Lesson Plan and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
from data_domain5 import DOMAIN5
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4+DOMAIN5
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

# ---------------- block DSL (single content stream → MD + DOCX) ----------------
B=[]
def h1(t): B.append(("h1",t))
def h2(t): B.append(("h2",t))
def h3(t): B.append(("h3",t))
def p(t):  B.append(("p",t))
def bullets(xs): B.append(("bullets",xs))
def steps(xs): B.append(("steps",xs))
def code(t): B.append(("code",t))
def note(t): B.append(("note",t))
def rule(): B.append(("rule",))

# ---------------- content ----------------
h1("Introduction")
p(f"This Learner Guide accompanies the WSQ course {C.TITLE} ({C.COURSE_CODE}), conducted by {C.ORG}. "
  "It provides step-by-step instructions for the 10 hands-on Autodesk Revit labs, organised by the four "
  f"course topics, and is aligned to the Skills Framework TSC {C.TSC_TITLE} ({C.TSC_CODE}). "
  "Every lab maps to a course learning outcome, and everything assessed in the final Written Assessment "
  "(K1\u2013K10) and Practical Performance (A1\u2013A9) is taught in these pages.")
p("Use this guide alongside the course slides and the lab dataset files in the labs/ folder of the course "
  "repository (also downloadable from the LMS). Each lab has its own folder containing a README with the "
  "full steps plus the Revit project files (.rvt) it uses. The guide is your open-book companion during "
  "the final assessment.")

h1("Course Learning Outcomes")
bullets(C.LEARNING_OUTCOMES)

h1("Skills Framework Alignment")
p(f"This course maps to the Skills Framework for the Built Environment TSC: {C.TSC_TITLE} ({C.TSC_CODE}).")
h3("TSC Knowledge (assessed by the Written Assessment)")
bullets([f"{k} \u2014 {v}" for k,v in C.TSC_KNOWLEDGE])
h3("TSC Abilities (assessed by the Practical Performance)")
bullets([f"{a} \u2014 {v}" for a,v in C.TSC_ABILITIES])

h1("Before You Start — Environment Setup")
h3("What you need")
bullets([
 "Autodesk Revit 2024 or 2025 installed on Windows (free 30-day trial, or an education licence from students.autodesk.com).",
 "The Revit sample projects installed with Revit (e.g. rac_basic_sample_project.rvt) for Labs 1 and 8.",
 "The course lab dataset files (.rvt, .pdf, .png) — in each lab's folder in the course repository, or downloadable from the LMS at https://lms-tms.tertiaryinfotech.com.",
 "A 3-button mouse (the middle button/wheel drives pan, zoom and orbit in Revit).",
])
h3("Launch and verify Revit")
p("Start Revit and confirm you can open the basic sample architecture project from the Home screen. "
  "Check the ribbon shows the Architecture, Structure, Insert, Annotate, Analyze, Massing & Site, Collaborate, View and Manage tabs — the labs use all of them.")
h3("Conventions used in every lab")
bullets([
 "Ribbon paths are written as Tab \u25b8 Panel \u25b8 Tool (e.g. Architecture \u25b8 Build \u25b8 Wall).",
 "Each lab folder contains the starter .rvt file(s) it needs; solution files (where provided) let you check your result.",
 "Save your work as <LabNN>_<YourName>.rvt so the trainer can review it.",
 "If a family or template is missing, load it via Insert \u25b8 Load Family or pick the closest available type.",
 "Imperial vs metric: the lab files use the units they were authored in — follow the values printed in the steps.",
])

# ---------------- per-topic, per-lab ----------------
TOPICS_BY_NUM={t["num"]:t for t in C.TOPICS}
for t in C.TOPICS:
    h1(f"Topic {t['code']} — {t['title']}  ({t['weighting']})")
    p(t["subtitle"])
    h3("Key concepts")
    bullets([(f"{x[0]} \u2014 {x[1]}" if isinstance(x,tuple) else x) for x in t["concepts"]])
    for a in [x for x in ACT if x["topic"]==t["num"]]:
        h2(f"Lab {a['num']} — {a['title']}")
        p(f"Objective: {a['objective']}.")
        p(f"Goal: {a['desc']}")
        h3("What you'll build")
        p(a["build"]+f"   (Tools: {a['services']}.)")
        h3("Step-by-step")
        st=[]
        for i,(instr,cmd) in enumerate(a["steps"],1):
            st.append((instr,cmd))
        steps(st)
        h3("Test it")
        p(a["test"])
        note(f"The full lab, with its dataset files, is in labs/lab-{a['num']:02d}-*/ in the course repository.")
        rule()

h1("Assessment Focus — How the Final Assessment Maps to This Guide")
p("The final assessment on Day 2 has two open-book instruments. Everything they assess is taught in this guide:")
h3("Written Assessment (SAQ) — 1 hour — assesses TSC Knowledge K1–K10")
bullets([
 "K1 Principles of BIM and K4 Definition of BIM — Topic 1 (What is BIM, the BIM process, BIM dimensions).",
 "K2 Value proposition of BIM — Topic 1 (the seven advantages of BIM).",
 "K3 Requirements of BIM — Topic 1 (OIR, AIR, EIR, BEP, PIM, AIM).",
 "K5 Application of BIM — Topic 3 (construction management, simulation, monitoring).",
 "K6 Technology used in BIM and K7 BIM design processes — Topics 2–3 (parametric modeling, Dynamo, analysis tools).",
 "K8 Documentation required for BIM — Topic 4 (BIM standard, BEP, BIP, PEP, EIR, MIDP, CDE, deliverables).",
 "K9 Databases and information systems — Topic 4 (SVY21/SHD, naming formats, file structures, worksharing).",
 "K10 BIM e-submission requirements and standards — Topic 4 (last saved views, cover page, core information, colour standards).",
])
h3("Practical Performance (PP) — 1.5 hours — assesses TSC Abilities A1–A9")
bullets([
 "A1–A2 (identify BIM application and interoperability) — practised in Labs 1–2 and the PP's project set-up tasks.",
 "A4–A5 (develop models, integrate design) — practised in Labs 2–6 (walls, floors, roofs, structure, linked models).",
 "A6–A8 (operate BIM, interpret output, analyse performance) — practised in Labs 7–8 (schedules, energy settings, checks).",
 "A3 and A9 (maintain databases, develop documentation) — practised in Labs 9–10 (worksharing, sheets, e-submission).",
 "The PP uses the provided PP dataset (.rvt); take a snapshot at the end of each task and paste it into the answer document.",
])
rule()

h1("Assessment Preparation")
bullets([
 "First pass: complete every lab in Revit, checking each lab's Test-it criterion.",
 "Second pass: redo the labs from memory until the ribbon paths and workflows are automatic.",
 "Review the Key Concepts of each topic — the WA questions are scenario-based versions of them.",
 "Practise the PP workflow: link a model, Copy/Monitor levels, configure Energy Settings, and export a documented report.",
 "Sharpen readiness with the Tertiary Infotech practice exams portal: https://exams.tertiaryinfotech.com.",
 "The assessment is open book: bring this guide and the slides, and know your way around them quickly.",
])

h1("Glossary")
gl=[
 ("BIM (Building Information Modelling)","An intelligent 3D model-based process for planning, designing, constructing and managing buildings and infrastructure."),
 ("BEP (BIM Execution Plan)","The Employer-approved baseline document defining goals, roles, deliverables and processes for BIM on a project."),
 ("OIR / AIR / EIR","Organisational, Asset and Employer Information Requirements — the chain of information needs BIM must satisfy."),
 ("PIM / AIM","Project Information Model (design & construction) / Asset Information Model (operation & asset management)."),
 ("CDE (Common Data Environment)","The single shared place where project information is stored, managed and exchanged."),
 ("MIDP","Master Information Delivery Plan — schedules who delivers which information, when."),
 ("LOD (Level of Development)","How developed a model element's geometry and information are — defined at project start."),
 ("IFC (Industry Foundation Classes)","The open, vendor-neutral exchange format BIM tools use for interoperability."),
 ("Parametric element","A Revit component whose behaviour is driven by parameters — one change updates every related view and schedule."),
 ("Family (system / loadable / in-place)","Revit's building blocks: predefined system elements, loadable RFA components, and unique in-project elements."),
 ("Toposurface / building pad","Terrain modelled from elevation points, and the level cut into it that the building sits on."),
 ("Worksharing / central model","Multiple team members editing local copies synchronised to one central model, divided into worksets."),
 ("Copy/Monitor","The Collaborate tool that copies levels/grids from a linked model and warns when the source changes."),
 ("Interference (clash) check","Revit's built-in detection of elements that occupy the same space — run before construction."),
 ("Quantity takeoff","Extracting element counts, materials and areas from the model for cost estimating."),
 ("SVY21 / SHD","The Singapore coordinate system and height datum a site model must be geo-referenced to for e-submission."),
 ("BIM e-submission","Submitting the BIM model to Singapore's regulatory agencies per the BCA Code of Practice — last saved views, cover page and core information."),
]
B.append(("dl",gl))

# ---------------- render Markdown ----------------
def _anchor(txt):
    return "".join(ch.lower() if ch.isalnum() else ("-" if ch in " -" else "") for ch in txt)

def render_md():
    out=[f"# {C.TITLE} — Learner Guide",""]
    out.append(f"**WSQ Course Code:** {C.COURSE_CODE}  |  **Conducted by:** {C.ORG} ({C.UEN.replace('UEN: ','UEN ')})  |  **Version {C.VERSION} · {C.VERSION_DATE}**")
    out.append("")
    # TOC (h1 + h2)
    out.append("## Contents"); out.append("")
    for kind,*rest in B:
        if kind=="h1": out.append(f"- [{rest[0]}](#{_anchor(rest[0])})")
        elif kind=="h2": out.append(f"  - [{rest[0]}](#{_anchor(rest[0])})")
    out.append("")
    for kind,*rest in B:
        if kind=="h1": out+=["",f"## {rest[0]}",""]
        elif kind=="h2": out+=["",f"### {rest[0]}",""]
        elif kind=="h3": out+=[f"**{rest[0]}**",""]
        elif kind=="p": out+=[rest[0],""]
        elif kind=="bullets": out+=[f"- {x}" for x in rest[0]]+[""]
        elif kind=="steps":
            for i,(instr,cmd) in enumerate(rest[0],1):
                out.append(f"{i}. {instr}")
                if cmd: out+=["",f"   ```bash",f"   {cmd}","   ```",""]
            out.append("")
        elif kind=="code": out+=["```bash",rest[0],"```",""]
        elif kind=="note": out+=[f"> **Note:** {rest[0]}",""]
        elif kind=="rule": out+=["---",""]
        elif kind=="dl":
            for term,defn in rest[0]: out.append(f"- **{term}** — {defn}")
            out.append("")
    return "\n".join(out)

MD_OUT=os.path.join(REPO,f"LG-{C.SHORT_TITLE}.md")
with open(MD_OUT,"w") as f: f.write(render_md())
print("Saved",MD_OUT)

# ---------------- render DOCX ----------------
BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
INKCODE=RGBColor(0x0B,0x30,0x60)
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc,"LEARNER GUIDE",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("1.0","23 May 2023","First version.","Tertiary Infotech Academy"),
 ("2.0","17 Oct 2025","Update company name.","Tertiary Infotech Academy"),
 (C.VERSION.lstrip("v"),C.VERSION_DATE,"Single-source rebuild aligned with the v10 slide deck and Lesson Plan — step-by-step guides for the 10 hands-on Revit labs, Skills Framework K/A alignment, assessment-focus mapping and a BIM glossary.",C.TRAINER),
])
prodoc.add_toc(doc)

def code_para(text):
    for line in text.split("\n"):
        para=doc.add_paragraph(); prodoc._shade_para(para) if hasattr(prodoc,"_shade_para") else None
        r=para.add_run(line); r.font.name="Consolas"; r.font.size=Pt(9.5); r.font.color.rgb=INKCODE

for kind,*rest in B:
    if kind=="h1": doc.add_heading(rest[0],level=1)
    elif kind=="h2": doc.add_heading(rest[0],level=2)
    elif kind=="h3":
        para=doc.add_paragraph(); r=para.add_run(rest[0]); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=BRAND
    elif kind=="p": doc.add_paragraph(rest[0])
    elif kind=="bullets":
        for x in rest[0]: doc.add_paragraph(x,style="List Bullet")
    elif kind=="steps":
        for i,(instr,cmd) in enumerate(rest[0],1):
            para=doc.add_paragraph(style="List Number"); para.add_run(instr)
            if cmd: code_para(cmd)
    elif kind=="code": code_para(rest[0])
    elif kind=="note":
        para=doc.add_paragraph(); r=para.add_run("Note: "); r.bold=True; r.font.color.rgb=BRAND
        para.add_run(rest[0]).font.size=Pt(10)
    elif kind=="rule": doc.add_paragraph("")
    elif kind=="dl":
        for term,defn in rest[0]:
            para=doc.add_paragraph(style="List Bullet")
            r=para.add_run(term+" — "); r.bold=True; para.add_run(defn)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT=os.path.join(REPO,"courseware",f"LG-{C.SHORT_TITLE}.docx")
doc.save(DOCX_OUT)
print("Saved",DOCX_OUT)
