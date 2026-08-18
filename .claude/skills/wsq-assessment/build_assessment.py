#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build the WSQ assessment set for 'Application of BIM using Revit' (TGS-2022015370):
  - Written Assessment (SAQ)  — 10 open-ended KNOWLEDGE questions (K1–K10), aligned to the slides
  - Practical Performance (PP) — 4 PRACTICAL tasks (A1–A9), aligned to the in-class labs
Each instrument is produced as a Question Paper and a matching Answer Key (4 DOCX total),
all with the WSQ house cover page (same as the Lesson Plan / Learner Guide). Page 1 is the cover;
page 2 carries Trainee Information + Instructions + Grading; the questions/tasks begin on page 3.
Body: Arial 11.
"""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# This script lives in the wsq-assessment skill (.claude/skills/wsq-assessment/) and runs in
# place — it detects the course repo root by walking up to the nearest dir that has a .git
# folder (or both courseware/ and assessment/). Override with env REPO=/path if needed.
def _find_repo():
    env = os.environ.get("REPO")
    if env and os.path.isdir(env):
        return os.path.abspath(env)
    d = os.path.dirname(os.path.abspath(__file__))
    while d != os.path.dirname(d):
        if os.path.isdir(os.path.join(d, ".git")) or \
           (os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "assessment"))):
            return d
        d = os.path.dirname(d)
    return os.getcwd()

REPO = _find_repo()
# prodoc.py (WSQ cover page + version control + page numbers, same as LP/LG) ships with the
# tertiary-lesson-plan skill. Look for it at the project level first, then the user level.
for _cand in (os.path.join(REPO, ".claude/skills/tertiary-lesson-plan"),
              os.path.expanduser("~/.claude/skills/tertiary-lesson-plan")):
    if os.path.exists(os.path.join(_cand, "prodoc.py")):
        sys.path.insert(0, _cand); break
import prodoc  # cover page + version control + page numbers (same as LP/LG)

# ─── EDIT PER COURSE ────────────────────────────────────────────────────────
TITLE       = "Application of BIM using Revit"
COURSE_CODE = "TGS-2022015370"
# ────────────────────────────────────────────────────────────────────────────
# The cover page renders prodoc's module-level TGS constant. Override it so the
# assessment cover shows THIS course's ref (works with either prodoc version —
# the older project prodoc has no course_code kwarg).
prodoc.TGS = f"TGS Ref No: {COURSE_CODE}"
OUT   = os.path.join(REPO, "assessment")

# Logos: prefer the course's own courseware/assets, else fall back to the copies bundled
# in this skill (so the assessment builds even outside this project). Replace the course
# logo per course; the Tertiary Infotech logo is the same for every WSQ course.
def _logo(name):
    here = os.path.dirname(os.path.abspath(__file__))
    for p in (os.path.join(REPO, "courseware/assets", name), os.path.join(here, "assets", name)):
        if os.path.exists(p):
            return p
    return None
ORG_LOGO    = _logo("tertiary-infotech-logo.png")
COURSE_LOGO = None   # None if absent → Tertiary-only cover (as LP/LG)

Q_VER, A_VER = "v4", "v4"   # WA: original v3 -> v4.  PP uses PP_VER below (original v6 -> v7).
PP_VER = "v7"
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)
# Assessments carry the cover page only — no Document Version Control Record.

# ---------------------------------------------------------------- WRITTEN (KNOWLEDGE)
# (criterion, context, question, [model-answer points]) — each traces to the course slides.
WRITTEN = [
 ("K1",
  "You are the BIM manager of an architecture practice, onboarding a project team that is using "
  "Building Information Modelling for the first time on a new commercial development.",
  "Explain three key principles of BIM that the team must understand before modelling begins.",
  ["BIM is one shared, intelligent, data-rich model — every stakeholder works from a single coordinated "
   "source of truth rather than separate drawings.",
   "BIM is a lifecycle process — Plan, Design, Build, Operate — the model's data serves the asset from "
   "planning through operations and maintenance.",
   "BIM carries dimensions beyond geometry: 3D geometry, 4D time, 5D cost, 6D sustainability, "
   "7D facility management (and 8D safety).",
   "(Slides: Principles of BIM / The Process of BIM / BIM Dimensions.)"]),
 ("K2",
  "A property developer is hesitant to fund BIM on a condominium project because of the up-front "
  "software and training investment. You must convince them of the returns.",
  "Present two key value propositions of BIM that justify the investment to the developer.",
  ["Maximised efficiency — shorter project lifecycle; everyone works off one up-to-date model, "
   "reducing errors and rework.",
   "Reduced costs and wastage — better material choices, streamlined construction and fewer human errors.",
   "Improved cost estimates — model-based estimates and quantity takeoffs are more realistic and accurate.",
   "Better insight into the project — realistic 3D visualisation lets changes happen in pre-construction, "
   "not on site; plus closer communication/collaboration, less risk and better end results.",
   "(Slides: Value Proposition — Advantages of BIM 1–7.)"]),
 ("K3",
  "A main contractor is preparing to bid for a project whose contract mandates BIM deliverables, and "
  "must put the foundational information requirements in place first.",
  "Identify two essential information requirements the company must address for effective BIM adoption, "
  "and state the purpose of each.",
  ["OIR — Organisational Information Requirements: what the organisation needs to know to run its business.",
   "AIR — Asset Information Requirements: what it needs to know about the assets it is responsible for.",
   "EIR — Employer Information Requirements: the information to be delivered, and the standards and "
   "processes to be adopted, for the project.",
   "BEP — BIM Execution Plan: the overall vision and implementation details the project team follows; "
   "PIM/AIM — the project and asset information models that carry the data.",
   "(Slides: Requirements of BIM — the Information Chain / The BIM Execution Plan.)"]),
 ("K4",
  "You have been invited to brief polytechnic built-environment students who have only ever used "
  "2D CAD drafting tools.",
  "How would you define BIM to the students, emphasising its role in modern construction practice?",
  ["Building Information Modelling (BIM) is an intelligent 3D model-based process that gives AEC "
   "professionals every detail needed to plan, design, construct and manage buildings and infrastructure.",
   "It is a process of sharing information — teams capture data as they work, and that data benefits "
   "operations, maintenance and future planning.",
   "It spans architecture, civil engineering, construction, plant, MEP and structural engineering, "
   "improving design decisions and building performance across the lifecycle.",
   "(Slides: What is BIM? / BIM in the Building Lifecycle.)"]),
 ("K5",
  "A project director wants to use BIM beyond design — to raise quality and efficiency during the "
  "construction phase of a new residential project.",
  "Briefly explain three applications of BIM the project director could use to improve project "
  "delivery and outcomes.",
  ["Clash detection / coordination — interference checks on the shared model resolve conflicts between "
   "disciplines before they reach site.",
   "Construction simulation and real-time monitoring — process management becomes visual and controllable, "
   "strengthening quality management.",
   "Quantity takeoff and cost estimation straight from the model; plus solar/energy analysis, logistics "
   "and facilities management.",
   "(Slides: Application of BIM / Application of BIM in Construction Management.)"]),
 ("K6",
  "Your innovation team is evaluating which technologies to pair with the existing BIM workflow to "
  "improve project outcomes.",
  "Name three technologies commonly used with BIM and explain how each supports the BIM process.",
  ["Cloud collaboration — up-to-date models accessible to every stakeholder, anywhere, any time.",
   "Drones / reality capture — real-world context data feeds the Plan phase of the model.",
   "AR/VR — immersive design review and client communication from the 3D model.",
   "IoT sensors — live asset data flowing into the 7D facility-management model; 3D printing and "
   "generative design (Dynamo) extend fabrication and design exploration.",
   "(Slides: Technologies Used in BIM / Generative Design and Dynamo.)"]),
 ("K7",
  "A design firm is restructuring its workflow to move from traditional 2D CAD drafting to a BIM "
  "design process.",
  "Briefly explain the BIM design process and how it differs from the traditional design process.",
  ["The BIM process runs Plan → Design → Build → Operate on one intelligent model: design authoring, "
   "analysis, detailing and documentation all happen in the parametric model.",
   "One change updates every related view, schedule and sheet automatically — no repetitive redrafting.",
   "Traditional 2D design is sequential and drawing-centric: each drawing is separate, coordination is "
   "manual, and changes must be chased through every sheet.",
   "BIM enables concurrent, collaborative working — shared models reduce conflicts between disciplines.",
   "(Slides: The Process of BIM / Develop Models for Building Elements.)"]),
 ("K8",
  "An architectural firm is preparing a BIM-based design proposal for submission to a regulatory "
  "body, and must assemble the governing documents first.",
  "List the documents required for a BIM project submission and state the purpose of each.",
  ["BIM Standard — the standards the project's BIM must follow; BEP — BIM Execution Plan, the "
   "Employer-approved baseline for goals, roles and deliverables.",
   "BIP — BIM Implementation Plan; PEP — Project Execution Plan.",
   "EIR — Employer Information Requirements; MIDP — Master Information Delivery Plan scheduling deliverables.",
   "CDE — Common Data Environment where all project information is shared; plus the agreed BIM "
   "deliverables (site/massing/discipline models, schedules, shop drawings, as-builts, FM data).",
   "(Slides: Documentation Required for BIM / BIM Deliverables.)"]),
 ("K9",
  "A construction company is expanding its digital infrastructure to support BIM projects delivered "
  "in Singapore.",
  "Describe the databases and information systems required for efficient BIM implementation.",
  ["A central model with worksharing — worksets let multiple team members edit local copies "
   "synchronised to one coordinated central model.",
   "Single or federated/linked file structures (UNC or relative paths; folders no more than two "
   "sub-folders deep) for multi-building or large projects.",
   "Geo-referencing to Singapore standards: SVY21 coordinates (x, y) and Singapore Height Datum (z); "
   "models at full 1:1 metric scale.",
   "Disciplined file and view naming formats and a Common Data Environment so the database stays "
   "navigable for every party.",
   "(Slides: Databases and Information Systems for BIM / Model Orientation, Site Configuration and Scale.)"]),
 ("K10",
  "You are responsible for preparing the BIM e-submission of a new commercial building project to "
  "Singapore's regulatory agencies.",
  "List 5 key BIM e-submission documentation requirements and standards that must be adhered to.",
  ["Last Saved Views — each view saved at maximum extent, no hidden objects/annotations, all external "
   "files loadable, irrelevant layers and drafting work purged, legible non-proprietary fonts.",
   "Cover page — submission authority and endorsements / QP's declaration, project information, and the "
   "list of views, schedules and sheets for approval.",
   "Core Information (CI) — the minimal data set specified by the respective regulatory agency.",
   "Colour standards — prescribed colours for amendments to approved plans and for A&A works.",
   "File/view naming per the Code of Practice for BIM e-Submission; 2D views (site plans, floor plans, "
   "elevations, sections) generated from the model at standard scales.",
   "(Slides: e-Submission — Last Saved Views Checklist / Cover Page and Core Information / Colour Standards.)"]),
]

# ---------------------------------------------------------------- PRACTICAL (ACTIVITY-BASED)
SCENARIO = (
 "You have been appointed as the BIM Coordinator for a new mixed-use development, \u201cHarbourview "
 "Exchange\u201d, comprising an office tower over a retail podium in a dense urban precinct. The client has "
 "commissioned a fully integrated BIM delivery: architectural, structural and MEP models must be "
 "coordinated in Revit, the building must meet energy-efficiency targets, and the project information "
 "must be maintained in a central database with documentation ready for stakeholder review and "
 "regulatory e-submission. For this assessment, use the provided \u201cApplication of BIM using Revit_PP.rvt\u201d "
 "dataset (download it from the LMS or get it from the assessor). You will perform tasks in Revit that "
 "demonstrate BIM application and interoperability, model integration, software operation, performance "
 "analysis and documentation development \u2014 each task mirrors a hands-on lab you completed in class.")

# (label, criterion, task prompt, box caption, model-answer build steps citing the lab)
PRACTICAL = [
 ("Task 1", "A1, A2",
  "Set up the mechanical model and demonstrate BIM application and interoperability. "
  "Open Revit and begin a new project using the Mechanical-Default template, and save it with an "
  "appropriate name for the Harbourview Exchange project. Open the Systems tab and navigate to the "
  "Mechanical Settings; adjust parameters relevant to a commercial project \u2014 duct sizing, airflow rates "
  "and pressure drops \u2014 based on industry standards, and save your changes. "
  "(Mirrors Lab 1 \u2014 Get Started with Revit; Lab 2 \u2014 project set-up.)",
  "Take a snapshot at the end of each point and paste the snapshots in the box below (A1, A2):",
  "1. Launch Revit \u25b8 File \u25b8 New \u25b8 Project \u25b8 Template: Mechanical-Default.rte \u25b8 OK.\n"
  "2. File \u25b8 Save As \u25b8 name the project (e.g. HVE_Mechanical_01.rvt) in the project folder.\n"
  "3. Systems tab \u25b8 Mechanical Settings (panel dialog launcher):\n"
  "   - Duct Sizing: set the default duct sizes for the commercial system.\n"
  "   - Airflow: configure supply/return airflow rates for the spaces.\n"
  "   - Pressure Drops: enter the expected pressure-drop values.\n"
  "4. Click OK to save the settings; snapshot each dialog as evidence.\n"
  "Evidence: new project saved from the Mechanical template with configured mechanical settings."),
 ("Task 2", "A4, A5",
  "Integrate the architectural design into your mechanical model. "
  "Use Link Revit to insert the provided \u201cApplication of BIM using Revit_PP.rvt\u201d architectural model "
  "into your project with positioning set to Auto \u2014 Origin to Origin. Confirm the linked model is "
  "correctly aligned, adjusting with Move or Rotate if necessary. Then use Copy/Monitor to copy the "
  "levels from the linked architectural model into your host mechanical model so that changes are "
  "monitored. (Mirrors Lab 9 \u2014 Sheets, Worksharing and Linked Models.)",
  "Take a snapshot at the end of each point and paste the snapshots in the box below (A4, A5):",
  "1. Insert \u25b8 Link \u25b8 Link Revit \u25b8 select Application of BIM using Revit_PP.rvt \u25b8 "
  "Positioning: Auto \u2014 Origin to Origin \u25b8 Open.\n"
  "2. Verify alignment against grids/walls; Modify \u25b8 Move / Rotate to adjust if required.\n"
  "3. Collaborate \u25b8 Coordinate \u25b8 Copy/Monitor \u25b8 Select Link \u25b8 pick the linked model.\n"
  "4. Copy \u25b8 pick each level in the linked model to copy it into the host; Finish.\n"
  "5. The copied levels are now monitored \u2014 Revit warns when the linked source changes.\n"
  "Evidence: linked architectural model correctly positioned; host levels copied and monitored."),
 ("Task 3", "A6, A7, A8",
  "Operate the analysis tools, interpret the output and analyse design performance. "
  "Go to the Analyze tab and open the Energy Settings dialog; input the project location, building "
  "type and HVAC system parameters, and save the settings. Create spaces/system zones for the "
  "functional areas, then run the heating and cooling loads analysis and review the results report. "
  "(Mirrors Lab 8 \u2014 Energy Analysis and Design Compliance Checks.)",
  "Take a snapshot at the end of each point and paste the snapshots in the box below (A6, A7, A8):",
  "1. Analyze \u25b8 Energy Analysis \u25b8 Energy Settings: set Location (project address), Building Type "
  "(e.g. Office) and the HVAC system; OK to save.\n"
  "2. Analyze \u25b8 Spaces & Zones \u25b8 Space: place spaces for offices, lobby and mechanical rooms; "
  "define system zones with their areas/volumes.\n"
  "3. Analyze \u25b8 Reports & Schedules \u25b8 Heating and Cooling Loads \u25b8 Calculate.\n"
  "4. Review the loads report (Project Browser \u25b8 Reports \u25b8 Analysis Reports); note peak heating/"
  "cooling demands \u2014 the values also appear on each analytical space's properties.\n"
  "Evidence: saved energy settings, defined zones and a generated loads/analysis report."),
 ("Task 4", "A3, A9",
  "Maintain the BIM database and develop the project documentation. "
  "Save your project with all configured settings, linked models and system zones properly organised. "
  "Develop a progress report that includes the snapshots of each completed task and a short "
  "description of every configuration made. Save the project file with the agreed name and submit "
  "the documentation with your report. (Mirrors Lab 9 \u2014 sheets/worksharing; Lab 10 \u2014 BIM "
  "e-Submission Standards Check.)",
  "Take a snapshot at the end of each point and paste the snapshots in the box below (A3, A9):",
  "1. Purge unused elements and confirm views are tidy (no hidden objects/annotations left behind) \u2014 "
  "the database stays clean and navigable.\n"
  "2. File \u25b8 Save As \u25b8 save the project as HVE_Mechanical_01.rvt (all settings, links and zones "
  "included).\n"
  "3. Compile the progress report: one section per task, with the snapshots pasted and a description "
  "of each configuration (template used, mechanical settings, link + Copy/Monitor, energy settings, "
  "zones, loads results).\n"
  "4. Submit the .rvt and the report document on the LMS.\n"
  "Evidence: an organised, saved model plus a documented progress report \u2014 the BIM database "
  "maintained (A3) and BIM documentation developed (A9)."),
]

# ---------------------------------------------------------------- doc helpers
def base_doc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    return doc

def para(doc, text, size=11, bold=False, italic=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if align is not None: p.alignment = align
    return p

def heading(doc, text, size=13):
    para(doc, text, size=size, bold=True, color=BRAND, after=6, before=8)

def answer_box(doc, lines=None, code=None, height_pt=90):
    """1x1 bordered box. `lines` → bullet-style model answer; `code` → monospace
    code/YAML/command block (indentation preserved); neither → empty answer space."""
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    if code:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in code.split("\n"):
            b = cell.add_paragraph(style=None)
            b.paragraph_format.space_after = Pt(0); b.paragraph_format.space_before = Pt(0)
            rr = b.add_run(ln if ln else " ")
            rr.font.name = "Consolas"; rr.font.size = Pt(9)
            rr._element.rPr.rFonts.set(qn('w:cs'), "Consolas")
            wt = rr._element.find(qn('w:t'))
            if wt is not None: wt.set(qn('xml:space'), 'preserve')
    elif lines:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in lines:
            b = cell.add_paragraph(style=None); b.paragraph_format.left_indent = Inches(0.15)
            rr = b.add_run("•  " + ln); rr.font.size = Pt(10.5)
    else:
        # empty answer space
        tr = t.rows[0]._tr
        trPr = tr.get_or_add_trPr(); trh = OxmlElement('w:trHeight')
        trh.set(qn('w:val'), str(int(height_pt*20))); trh.set(qn('w:hRule'), 'atLeast'); trPr.append(trh)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

FILL_GAP = 6    # extra space below each fill-in line (paired with double line spacing for writing room)

def candidate_block(doc):
    heading(doc, "Trainee Information")
    for label in ["Trainee Name (as per NRIC): ______________________________________",
                  "Last 3 digits and alphabet of NRIC/FIN: ____________________",
                  "Date: ____________________"]:
        p = para(doc, label, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

# Assessment briefing (from the course slides — "Briefing for Assessment").
BRIEFING = [
    "Place phones and other materials under the table or on the floor.",
    "No photos or recording of assessment scripts.",
    "No discussion during the assessment.",
    "Use a black/blue pen for hard-copy assessments.",
    "No liquid paper / correction tape.",
    "Scripts are collected when time is up.",
]

LMS_URL = "https://lms-tms.tertiaryinfotech.com/"

def add_hyperlink(p, url, text):
    """Add a real clickable Word hyperlink (blue, underlined) to paragraph p."""
    r_id = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)  # 11pt
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run); p._p.append(link)
    return link

def instructions(doc, minutes_text):
    heading(doc, "Instructions to Candidate")
    # None marks the upload instruction, which carries a clickable LMS hyperlink.
    items = [
        "This is an individual exercise.",
        "This is an open-book assessment.",
        f"A total of {minutes_text} is given to complete this assessment.",
        None,
    ] + BRIEFING
    for i, s in enumerate(items, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        if s is None:
            p.add_run(f"{i}.  Complete your answers on the document provided and "
                      "upload the completed answers to the LMS at ").font.size = Pt(11)
            add_hyperlink(p, LMS_URL, LMS_URL)
            p.add_run(".").font.size = Pt(11)
        else:
            p.add_run(f"{i}.  {s}").font.size = Pt(11)

def grading(doc, what):
    heading(doc, "Grading")
    para(doc, what, size=11, after=12)
    for ln in ["Grade: _______  (C / NYC)",
               "Assessor Name: __________________________   Assessor NRIC: ________________",
               "Date: ________________________                    Signature: ____________________"]:
        p = para(doc, ln, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

def finish(doc, path):
    prodoc.add_page_numbers(doc); prodoc.enable_update_fields(doc)
    doc.save(path); print("  saved:", os.path.basename(path))

# ---------------------------------------------------------------- builders
def build_wa(answers):
    doc = base_doc()
    kind = "Written Assessment (SAQ) — Answer Key" if answers else "Written Assessment (SAQ)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Written Assessment (SAQ)" if answers else "Written Assessment (SAQ)",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; questions begin on the next page.
        candidate_block(doc); instructions(doc, "1 hour")
        grading(doc, "Candidate has answered all written questions and demonstrated the underpinning "
                     "knowledge required for the course learning outcomes.")
        page_break(doc)
    para(doc, "Short-Answer Questions (Knowledge)", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Answer all questions in your own words. Each question tests underpinning knowledge covered in the "
              "course slides.", size=10.5, italic=True, color=GREY, after=8)
    # Pagination is EXPLICIT — two questions to a page on the paper, one model answer to a
    # page in the key. Do not swap this for Word's keepNext/cantSplit: Word pushes an
    # oversized box to the next page, but Google Docs draws the border anyway and prints the
    # question text and the page footer straight THROUGH it. See SKILL.md → Pagination.
    per_page = 1 if answers else 2
    for i, (crit, ctx, q, pts) in enumerate(WRITTEN, 1):
        para(doc, f"Question {i}:", size=11.5, bold=True, after=2, before=6)
        para(doc, ctx, size=11, after=3)
        para(doc, f"{q}  ({crit})", size=11, bold=True, after=4)
        answer_box(doc, lines=pts if answers else None)
        if i % per_page == 0 and i < len(WRITTEN):
            page_break(doc)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to WA (SAQ) - {TITLE} - {suffix}.docx" if answers
            else f"WA (SAQ) - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

def build_pp(answers):
    doc = base_doc()
    kind = "Practical Performance (PP) — Answer Key" if answers else "Practical Performance (PP)"
    prodoc.add_cover_page(doc, kind, TITLE, PP_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Practical Performance Assessment" if answers else "Practical Performance Assessment",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; the problem begins on the next page.
        candidate_block(doc); instructions(doc, "90 minutes")
        grading(doc, "Candidate has successfully completed all PP tasks and can explain the overall "
                     "functions and features used to achieve them.")
        page_break(doc)
    para(doc, "Practical Problem", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Scenario", size=11.5, bold=True, after=2)
    para(doc, SCENARIO, size=11, after=8)
    # Practical tasks are long and their boxes are tall, so they get a page each — on the
    # paper AND in the key. Same rule as the WA: the page break is ours, not the renderer's.
    for i, (label, crit, prompt, cap, pts) in enumerate(PRACTICAL, 1):
        para(doc, f"{label} ({crit}):", size=11.5, bold=True, after=2, before=6)
        para(doc, prompt, size=11, after=3)
        para(doc, cap, size=10.5, italic=True, color=GREY, after=4)
        answer_box(doc, code=pts if answers else None, height_pt=150)
        if i < len(PRACTICAL):
            page_break(doc)
    suffix = PP_VER
    name = (f"Answer to PP Assessment - {TITLE} - {suffix}.docx" if answers
            else f"PP Assessment - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

if __name__ == "__main__":
    print("Building WSQ assessment set…")
    build_wa(answers=False); build_wa(answers=True)
    build_pp(answers=False); build_pp(answers=True)
    print(f"Done. WA: {len(WRITTEN)} questions · PP: {len(PRACTICAL)} tasks.")
